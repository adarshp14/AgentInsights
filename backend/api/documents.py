"""
Multi-tenant document management API
"""
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from sqlalchemy.orm import Session
from typing import List, Dict, Any
import uuid
import hashlib
import os
from datetime import datetime

from database.database import get_db
from database.models import Document, User
from services.auth_service import AuthService
from schemas.auth import UserProfile
from auth.clerk_auth import get_current_user_compatible
from retriever.multi_tenant_vector_store import MultiTenantVectorStore, get_vector_store
from langchain.schema import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter

router = APIRouter(prefix="/documents", tags=["Document Management"])

# Use Clerk authentication or fallback to custom auth
def get_current_user(
    # Try Clerk auth first, fallback to custom auth
    user_info: Dict[str, Any] = Depends(get_current_user_compatible)
) -> Dict[str, Any]:
    """Get current authenticated user (Clerk or custom auth)"""
    return user_info

@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: Session = Depends(get_db),
    vector_store: MultiTenantVectorStore = Depends(get_vector_store)
):
    """
    Upload document to organization's knowledge base
    Requires authentication and admin role
    """
    # Check if user has admin permissions
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only admin users can upload documents"
        )
    
    try:
        # Validate file type
        allowed_types = ['pdf', 'txt', 'docx', 'text']
        file_extension = file.filename.split('.')[-1].lower() if '.' in file.filename else 'txt'
        
        if file_extension not in allowed_types:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_types)}"
            )
        
        # Read file content
        content = await file.read()
        file_size = len(content)
        
        # Create content hash for deduplication
        content_hash = hashlib.sha256(content).hexdigest()
        
        # Check if document already exists
        existing_doc = db.query(Document).filter(
            Document.org_id == current_user.org_id,
            Document.content_hash == content_hash
        ).first()
        
        if existing_doc:
            return {
                "status": "duplicate",
                "message": f"Document with same content already exists: {existing_doc.filename}",
                "existing_document_id": str(existing_doc.document_id)
            }
        
        # Generate unique filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{timestamp}_{file.filename}"
        
        # Create document record
        document = Document(
            org_id=current_user.org_id,
            filename=safe_filename,
            original_filename=file.filename,
            file_type=file_extension,
            file_size=file_size,
            file_path=f"uploads/{current_user.org_id}/{safe_filename}",
            processing_status="processing",
            uploaded_by=current_user.user_id,
            content_hash=content_hash,
            embedding_model="all-MiniLM-L6-v2",
            doc_metadata={
                "uploaded_via": "admin_api",
                "uploader_email": current_user.email
            }
        )
        
        db.add(document)
        db.flush()  # Get the document_id
        
        # Process document content
        text_content = ""
        if file_extension == 'txt' or file_extension == 'text':
            text_content = content.decode('utf-8')
        elif file_extension == 'pdf':
            try:
                import PyPDF2
                import io
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text_content = "\\n".join([page.extract_text() for page in pdf_reader.pages])
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Error processing PDF: {str(e)}"
                )
        elif file_extension == 'docx':
            try:
                from docx import Document as DocxDocument
                import io
                docx_doc = DocxDocument(io.BytesIO(content))
                text_content = "\\n".join([paragraph.text for paragraph in docx_doc.paragraphs])
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Error processing DOCX: {str(e)}"
                )
        
        # Update document with extracted text length
        document.extracted_text_length = len(text_content)
        
        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        
        chunks = text_splitter.split_text(text_content)
        
        # Create Langchain documents
        langchain_docs = [
            LangchainDocument(
                page_content=chunk,
                metadata={
                    "source": document.filename,
                    "document_id": str(document.document_id),
                    "chunk_index": i,
                    "file_type": file_extension,
                    "org_id": str(current_user.org_id)
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        
        # Add to vector store
        vector_result = vector_store.add_documents(
            org_id=current_user.org_id,
            documents=langchain_docs,
            document_id=document.document_id
        )
        
        if vector_result["success"]:
            document.processing_status = "completed"
            document.chunks_created = vector_result["chunks_added"]
            document.processed_date = datetime.utcnow()
        else:
            document.processing_status = "failed"
            
        db.commit()
        
        return {
            "status": "success",
            "document_id": str(document.document_id),
            "filename": document.filename,
            "file_type": file_extension,
            "file_size": file_size,
            "chunks_created": document.chunks_created,
            "processing_status": document.processing_status,
            "org_id": str(current_user.org_id)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error processing document: {str(e)}"
        )

@router.get("/list")
async def list_documents(
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all documents in organization"""
    try:
        documents = db.query(Document).filter(
            Document.org_id == current_user["org_id"]
        ).order_by(Document.upload_date.desc()).all()
        
        doc_list = []
        for doc in documents:
            doc_list.append({
                "document_id": str(doc.document_id),
                "filename": doc.original_filename,
                "file_type": doc.file_type,
                "file_size": doc.file_size,
                "chunks_created": doc.chunks_created,
                "processing_status": doc.processing_status,
                "upload_date": doc.upload_date.isoformat(),
                "processed_date": doc.processed_date.isoformat() if doc.processed_date else None,
                "uploader": doc.uploader.full_name if doc.uploader else "Unknown"
            })
        
        return {
            "documents": doc_list,
            "total": len(doc_list),
            "org_id": str(current_user.org_id)
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error listing documents: {str(e)}"
        )

@router.delete("/{document_id}")
async def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    vector_store: MultiTenantVectorStore = Depends(get_vector_store)
):
    """Delete document from organization's knowledge base"""
    # Check admin permissions
    if current_user.role != "admin":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only admin users can delete documents"
        )
    
    try:
        doc_uuid = uuid.UUID(document_id)
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID format"
        )
    
    try:
        # Find document
        document = db.query(Document).filter(
            Document.document_id == doc_uuid,
            Document.org_id == current_user.org_id
        ).first()
        
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        # Delete from vector store
        vector_result = vector_store.delete_document(
            org_id=current_user.org_id,
            document_id=doc_uuid
        )
        
        # Delete from database
        db.delete(document)
        db.commit()
        
        return {
            "status": "success",
            "message": f"Document '{document.original_filename}' deleted successfully",
            "chunks_deleted": vector_result.get("chunks_deleted", 0)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error deleting document: {str(e)}"
        )

@router.get("/stats")
async def get_document_stats(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    vector_store: MultiTenantVectorStore = Depends(get_vector_store)
):
    """Get document statistics for organization"""
    try:
        # Database stats
        total_docs = db.query(Document).filter(
            Document.org_id == current_user.org_id
        ).count()
        
        completed_docs = db.query(Document).filter(
            Document.org_id == current_user.org_id,
            Document.processing_status == "completed"
        ).count()
        
        total_size = db.query(db.func.sum(Document.file_size)).filter(
            Document.org_id == current_user.org_id
        ).scalar() or 0
        
        # Vector store stats
        vector_stats = vector_store.get_org_stats(current_user.org_id)
        
        return {
            "total_documents": total_docs,
            "completed_documents": completed_docs,
            "processing_documents": total_docs - completed_docs,
            "total_size_bytes": total_size,
            "total_size_mb": round(total_size / (1024 * 1024), 2),
            "total_chunks": vector_stats.get("total_chunks", 0),
            "vector_store_collection": vector_stats.get("collection_name", ""),
            "org_id": str(current_user.org_id)
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error getting document stats: {str(e)}"
        )